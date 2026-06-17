from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
MARKDOWN_PATH = ROOT / "HDB_Chatbot_Implementation_Guide.md"
OUTPUT_PATH = ROOT / "HDB_Chatbot_Walkthrough.docx"

TITLE = "HDB Chatbot Implementation Walkthrough"
SUBTITLE = "Local setup, architecture, operations, and demo runbook for the hdb-bot project"

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(34, 34, 34)
MUTED = RGBColor(90, 96, 110)
LIGHT_BG = "F3F5F7"
BORDER = "D7DBE2"


def set_run_font(run, name: str = "Arial", size: int | None = None,
                 color: RGBColor | None = None, bold: bool | None = None,
                 italic: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_paragraph(paragraph, fill: str = LIGHT_BG) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_bottom_rule(paragraph, color: str = BORDER, size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("HDB Bot Implementation Walkthrough")
    set_run_font(run, size=9, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("sai-harshita | hdb-bot")
    set_run_font(run, size=9, color=MUTED)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, NAVY),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(14 if style_name == "Heading 1" else 10)
        style.paragraph_format.space_after = Pt(6 if style_name != "Heading 3" else 4)


def add_title_block(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run(TITLE)
    set_run_font(run, size=23, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run(SUBTITLE)
    set_run_font(run, size=12, color=MUTED)

    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.autofit = False
    meta_table.columns[0].width = Inches(1.8)
    meta_table.columns[1].width = Inches(4.7)
    rows = [
        ("Repository", "https://github.com/sai-harshita/hdb-bot"),
        ("Deployment mode", "Local Docker Compose stack with optional public tunnel or VM migration"),
    ]
    for row, (label, value) in zip(meta_table.rows, rows):
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(4.7)
        p0 = row.cells[0].paragraphs[0]
        p1 = row.cells[1].paragraphs[0]
        p0.paragraph_format.space_after = Pt(2)
        p1.paragraph_format.space_after = Pt(2)
        r0 = p0.add_run(label)
        r1 = p1.add_run(value)
        set_run_font(r0, size=10, color=DARK, bold=True)
        set_run_font(r1, size=10, color=DARK)

    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(8)
    after.paragraph_format.space_after = Pt(8)
    add_bottom_rule(after)


def flush_paragraph_buffer(doc: Document, buffer: list[str]) -> None:
    if not buffer:
        return
    paragraph = doc.add_paragraph(" ".join(part.strip() for part in buffer if part.strip()))
    paragraph.paragraph_format.space_after = Pt(6)
    buffer.clear()


def add_code_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.left_indent = Inches(0.15)
    paragraph.paragraph_format.right_indent = Inches(0.15)
    shade_paragraph(paragraph)
    run = paragraph.add_run(text.rstrip())
    set_run_font(run, name="Consolas", size=9, color=DARK)


def build_document() -> None:
    markdown = MARKDOWN_PATH.read_text(encoding="utf-8")
    lines = markdown.splitlines()

    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    add_title_block(doc)

    paragraph_buffer: list[str] = []
    in_code_block = False
    first_h1_consumed = False

    for raw_line in lines:
        line = raw_line.rstrip("\n")
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_paragraph_buffer(doc, paragraph_buffer)
            in_code_block = not in_code_block
            continue

        if in_code_block:
            add_code_paragraph(doc, raw_line)
            continue

        if not stripped:
            flush_paragraph_buffer(doc, paragraph_buffer)
            continue

        if stripped.startswith("# "):
            flush_paragraph_buffer(doc, paragraph_buffer)
            if not first_h1_consumed:
                first_h1_consumed = True
                continue
            doc.add_paragraph(stripped[2:].strip(), style="Heading 1")
            continue

        if stripped.startswith("## "):
            flush_paragraph_buffer(doc, paragraph_buffer)
            doc.add_paragraph(stripped[3:].strip(), style="Heading 1")
            continue

        if stripped.startswith("### "):
            flush_paragraph_buffer(doc, paragraph_buffer)
            doc.add_paragraph(stripped[4:].strip(), style="Heading 2")
            continue

        if re.match(r"^\d+\.\s", stripped):
            flush_paragraph_buffer(doc, paragraph_buffer)
            paragraph = doc.add_paragraph(style="List Number")
            run = paragraph.add_run(re.sub(r"^\d+\.\s", "", stripped))
            set_run_font(run, size=11, color=DARK)
            continue

        if stripped.startswith("- "):
            flush_paragraph_buffer(doc, paragraph_buffer)
            paragraph = doc.add_paragraph(style="List Bullet")
            run = paragraph.add_run(stripped[2:].strip())
            set_run_font(run, size=11, color=DARK)
            continue

        paragraph_buffer.append(stripped)

    flush_paragraph_buffer(doc, paragraph_buffer)
    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
